import os
import re
import json
import boto3
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor

load_dotenv()

REGION = os.getenv("AWS_DEFAULT_REGION", "us-east-1")

# Initialize AWS clients
bedrock_runtime = boto3.client("bedrock-runtime", region_name=REGION)
s3_client = boto3.client("s3", region_name=REGION)
lambda_client = boto3.client("lambda", region_name=REGION)
iam_client = boto3.client("iam", region_name=REGION)
dynamodb_client = boto3.client("dynamodb", region_name=REGION)
cloudwatch_client = boto3.client("cloudwatch", region_name=REGION)
amplify_client = boto3.client("amplify", region_name=REGION)

def collect_aws_data():
    print("Collecting live AWS data for ISO 27001 audit...")
    s3_data, lambda_data, iam_data, dynamo_data, cw_data, amp_data = [], [], [], [], [], []
    
    try:
        for b in s3_client.list_buckets().get("Buckets", [])[:5]:
            s3_data.append({"Bucket": b["Name"]})
    except Exception as e: pass
    
    try:
        for f in lambda_client.list_functions().get("Functions", [])[:5]:
            lambda_data.append({"Function": f["FunctionName"]})
    except Exception as e: pass
    
    try:
        for u in iam_client.list_users().get("Users", [])[:5]:
            mfa = len(iam_client.list_mfa_devices(UserName=u["UserName"]).get("MFADevices", [])) > 0
            iam_data.append({"User": u["UserName"], "MFA_Active": mfa})
    except Exception as e: pass
    
    try:
        for t in dynamodb_client.list_tables().get("TableNames", [])[:5]:
            dynamo_data.append({"Table": t})
    except Exception as e: pass
    
    try:
        cw_data.append({"Alarms": len(cloudwatch_client.describe_alarms().get("MetricAlarms", []))})
    except Exception as e: pass
    
    try:
        amp_data = [{"App": a["name"]} for a in amplify_client.list_apps().get("apps", [])]
    except Exception as e: pass
        
    return {"S3": s3_data, "Lambda": lambda_data, "IAM": iam_data, "DynamoDB": dynamo_data, "CloudWatch": cw_data, "Amplify": amp_data}

ISO_SYSTEM_PROMPT = """
You are an Information Security Auditor evaluating AWS configurations against ISO/IEC 27001:2022 Annex A controls.
Map AWS infrastructure configs to:
A.8.7 Protection against malware / config management
A.8.12 Data leakage prevention / S3 Public Access
A.8.24 Use of cryptography / Encryption at Rest
A.8.5 Secure authentication / IAM MFA
A.8.15 Logging and monitoring / CloudWatch

Format output strictly as follows:

# Executive Summary
Summary of ISO 27001 alignment.

# Security Score & Compliance Rating
Overall ISO 27001 Compliance Score: [X]/100
Compliant Controls: [X]
Non-Compliant Controls: [Y]

# Detailed ISO 27001 Annex A Findings
List controls with [COMPLIANT] or [NON-COMPLIANT], mapped ISO control ID, and remediation.
Use **bold text** for emphasis and status tags like **[COMPLIANT]** or **[NON-COMPLIANT]**.
"""

def run_iso_audit(aws_config):
    response = bedrock_runtime.converse(
        modelId="amazon.nova-lite-v1:0",
        messages=[{"role": "user", "content": [{"text": f"AWS Config:\n{json.dumps(aws_config, indent=2)}"}]}],
        system=[{"text": ISO_SYSTEM_PROMPT}]
    )
    return response['output']['message']['content'][0]['text']

def add_formatted_paragraph(doc, text, style=None):
    """Parses bold markdown tags and applies native Word bold formatting."""
    p = doc.add_paragraph(style=style)
    # Split text by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def export_to_word(audit_text, filename="AWS_ISO27001_Security_Audit_Report.docx"):
    doc = Document()
    
    # Uniform Title Styling (Matching NIST)
    title = doc.add_heading("AWS ISO/IEC 27001:2022 Security Audit Report", level=0)
    title.style.font.color.rgb = RGBColor(0, 51, 102)
    
    for line in audit_text.split("\n"):
        if line.startswith("# "):
            h = doc.add_heading(line.replace("# ", ""), level=1)
            h.style.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith("## "):
            h = doc.add_heading(line.replace("## ", ""), level=2)
            h.style.font.color.rgb = RGBColor(102, 102, 102)
        elif line.startswith("- ") or line.startswith("* "):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')
        elif line.strip():
            add_formatted_paragraph(doc, line)
            
    doc.save(filename)
    print(f"ISO 27001 Report saved to: {filename}")

if __name__ == "__main__":
    config = collect_aws_data()
    report = run_iso_audit(config)
    print(report)
    export_to_word(report)