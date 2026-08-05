import os
import re
import json
import boto3
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor

load_dotenv()
REGION = os.getenv("AWS_DEFAULT_REGION", "us-east-1")

bedrock_runtime = boto3.client("bedrock-runtime", region_name=REGION)
s3_client = boto3.client("s3", region_name=REGION)
lambda_client = boto3.client("lambda", region_name=REGION)
iam_client = boto3.client("iam", region_name=REGION)
dynamodb_client = boto3.client("dynamodb", region_name=REGION)
cloudwatch_client = boto3.client("cloudwatch", region_name=REGION)
amplify_client = boto3.client("amplify", region_name=REGION)

def collect_all_aws_config():
    print("Collecting live AWS data...")
    # Gather configuration data across services
    s3_data, lambda_data, iam_data, dynamo_data, cw_data, amp_data = [], [], [], [], [], []
    
    try:
        for b in s3_client.list_buckets().get("Buckets", [])[:5]:
            name = b["Name"]
            pab = s3_client.get_public_access_block(Bucket=name)["PublicAccessBlockConfiguration"] if True else "Not Configured"
            enc = s3_client.get_bucket_encryption(Bucket=name)["ServerSideEncryptionConfiguration"] if True else "None"
            ver = s3_client.get_bucket_versioning(Bucket=name).get("Status", "Disabled")
            s3_data.append({"Bucket": name, "PAB": pab, "Encryption": enc, "Versioning": ver})
    except: pass

    try:
        for f in lambda_client.list_functions().get("Functions", [])[:5]:
            lambda_data.append({"Function": f["FunctionName"], "VPC": bool(f.get("VpcConfig", {}).get("VpcId"))})
    except: pass

    try:
        for u in iam_client.list_users().get("Users", [])[:5]:
            mfa = len(iam_client.list_mfa_devices(UserName=u["UserName"]).get("MFADevices", [])) > 0
            iam_data.append({"User": u["UserName"], "MFA_Active": mfa})
    except: pass

    try:
        for t in dynamodb_client.list_tables().get("TableNames", [])[:5]:
            desc = dynamodb_client.describe_table(TableName=t)["Table"]
            dynamo_data.append({"Table": t, "SSE": desc.get("SSEDescription", {}).get("Status", "DISABLED")})
    except: pass

    try:
        alarms = len(cloudwatch_client.describe_alarms().get("MetricAlarms", []))
        cw_data.append({"MetricAlarmsCount": alarms})
    except: pass

    try:
        apps = amplify_client.list_apps().get("apps", [])
        amp_data = [{"App": a["name"]} for a in apps]
    except: pass

    return {"S3": s3_data, "Lambda": lambda_data, "IAM": iam_data, "DynamoDB": dynamo_data, "CloudWatch": cw_data, "Amplify": amp_data}

NIST_SYSTEM_PROMPT = """
You are a Cloud Security Auditor evaluating AWS against NIST SP 800-53 controls.
Examine S3, Lambda, IAM, DynamoDB, CloudWatch, and Amplify configuration data.

Return report strictly formatted as follows:

# Executive Summary
Provide a 2-paragraph overview.

# Security Score
Overall Score: [X]/100
Passed Controls: [X]
Failed Controls: [Y]

# Detailed Audit Findings
Format findings as list with status, control mapping, and remediation steps.
Use **bold text** for key findings and status tags like **[PASS]** or **[FAIL]**.
"""

def run_audit(aws_config):
    response = bedrock_runtime.converse(
        modelId="amazon.nova-lite-v1:0",
        messages=[{"role": "user", "content": [{"text": f"AWS Config:\n{json.dumps(aws_config, indent=2)}"}]}],
        system=[{"text": NIST_SYSTEM_PROMPT}]
    )
    return response['output']['message']['content'][0]['text']

def add_formatted_paragraph(doc, text, style=None):
    """Parses **bold** markdown tags and applies native Word bold formatting."""
    p = doc.add_paragraph(style=style)
    # Split text by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def export_to_word(audit_text, filename="AWS_NIST_Security_Audit_Report.docx"):
    doc = Document()
    title = doc.add_heading("AWS NIST SP 800-53 Security Audit Report", level=0)
    title.style.font.color.rgb = RGBColor(0, 51, 102)

    for line in audit_text.split("\n"):
        if line.startswith("# "):
            h = doc.add_heading(line.replace("# ", ""), level=1)
            h.style.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith("## "):
            h = doc.add_heading(line.replace("## ", ""), level=2)
            h.style.font.color.rgb = RGBColor(102, 102, 102)
        elif line.startswith("- ") or line.startswith("* "):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')
        elif line.strip():
            add_formatted_paragraph(doc, line)

    doc.save(filename)
    print(f"Report saved to: {filename}")

if __name__ == "__main__":
    config = collect_all_aws_config()
    report = run_audit(config)
    print(report)
    export_to_word(report)